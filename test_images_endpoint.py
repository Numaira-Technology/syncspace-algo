"""
End-to-end test for /run-syncspace-images/ endpoint.

Converts test DOCX and XLSX files to base64 PNG images and calls the endpoint.

Usage: python test_images_endpoint.py
Input: test_case/2_output_earnings_descr_2023Q2.docx, test_case/3_source_TSLA_2024Q2.xlsx
Output: JSON response from /run-syncspace-images/
"""
import base64
import json
import sys
import os
import requests
from PIL import Image, ImageDraw, ImageFont
import openpyxl
from docx import Document


DOCX_PATH = r"D:\Documents\Project Bird\code\syncspace-app\test_case\2_output_earnings_descr_2023Q2.docx"
XLSX_PATH = r"D:\Documents\Project Bird\code\syncspace-app\test_case\3_source_TSLA_2024Q2.xlsx"
API_URL = "http://localhost:8002/run-syncspace-images/"


def xlsx_to_base64_image(path: str) -> str:
    wb = openpyxl.load_workbook(path)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("Excel sheet is empty")

    col_count = max(len(r) for r in rows)
    col_widths = [80] * col_count
    for row in rows:
        for ci, cell in enumerate(row):
            text = str(cell) if cell is not None else ""
            w = min(300, max(80, int(len(text) * 7.5 + 16)))
            if w > col_widths[ci]:
                col_widths[ci] = w

    cell_h = 28
    total_w = sum(col_widths)
    total_h = len(rows) * cell_h

    img = Image.new("RGB", (total_w, total_h), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
        font_bold = ImageFont.truetype("arialbd.ttf", 12)
    except Exception:
        font = ImageFont.load_default()
        font_bold = font

    for ri, row in enumerate(rows):
        x = 0
        for ci, colw in enumerate(col_widths):
            cell_val = row[ci] if ci < len(row) else None
            text = str(cell_val) if cell_val is not None else ""
            y = ri * cell_h

            # Header row background
            if ri == 0:
                draw.rectangle([x, y, x + colw - 1, y + cell_h - 1], fill="#D0D7DE")

            # Border
            draw.rectangle([x, y, x + colw - 1, y + cell_h - 1], outline="#888888")

            # Truncate text
            max_w = colw - 16
            display_text = text
            while len(display_text) > 1 and draw.textlength(display_text, font=font) > max_w:
                display_text = display_text[:-1]
            if display_text != text:
                display_text += "…"

            f = font_bold if ri == 0 else font
            draw.text((x + 8, y + cell_h // 2 - 6), display_text, fill="black", font=f)
            x += colw

    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def docx_to_base64_image(path: str) -> str:
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append((text, para.style.name))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in [l[0] for l in lines]:
                    lines.append((text, "Normal"))

    width = 900
    line_h = 22
    padding = 20
    total_h = len(lines) * line_h + padding * 2

    img = Image.new("RGB", (width, total_h), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
        font_heading = ImageFont.truetype("arialbd.ttf", 14)
    except Exception:
        font = ImageFont.load_default()
        font_heading = font

    y = padding
    for text, style in lines:
        f = font_heading if "Heading" in style or "Title" in style else font
        color = "#333333" if "Heading" in style else "black"
        max_w = width - padding * 2
        display_text = text
        while len(display_text) > 1 and draw.textlength(display_text, font=f) > max_w:
            display_text = display_text[:-1]
        if display_text != text:
            display_text += "…"
        draw.text((padding, y), display_text, fill=color, font=f)
        y += line_h

    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def print_result(result: dict) -> bool:
    if result.get("status") == "success":
        data = result.get("data", result)
        if "results" not in data:
            # Nested: {"status": "success", "data": {"number_of_changes": ..., "results": [...]}}
            data = result.get("data", {})
        print(f"\nSUCCESS -- {data.get('number_of_changes', len(data.get('results', [])))} change(s) found")
        for i, r in enumerate(data.get("results", []), 1):
            print(f"\n  Change {i}:")
            print(f"    Original:  {r['original_text'][:120]}")
            print(f"    Modified:  {r['modified_text'][:120]}")
            print(f"    Confidence: {r['confidence']:.2f}")
        return True
    else:
        print(f"\nFAILED: {result.get('message', 'unknown error')}")
        return False


def main():
    print("Converting XLSX to image...")
    excel_b64 = xlsx_to_base64_image(XLSX_PATH)
    print(f"  Excel image: {len(excel_b64)} base64 chars")

    print("Converting DOCX to image...")
    docx_b64 = docx_to_base64_image(DOCX_PATH)
    print(f"  DOCX image: {len(docx_b64)} base64 chars")

    # Test 1: Direct algo service
    print(f"\n[Test 1] Calling algo directly: {API_URL}")
    payload_algo = {"docx_image": docx_b64, "excel_image": excel_b64}
    try:
        response = requests.post(API_URL, json=payload_algo, timeout=600)
        print(f"  HTTP status: {response.status_code}")
        if not print_result(response.json()):
            sys.exit(1)
    except Exception as e:
        print(f"  Algo request error: {e}")
        sys.exit(1)

    # Test 2: Through Flask backend
    backend_url = "http://localhost:8000/ai/newAiProcess"
    print(f"\n[Test 2] Calling Flask backend: {backend_url}")
    payload_backend = {"wordImage": docx_b64, "excelImage": excel_b64}
    try:
        response2 = requests.post(backend_url, json=payload_backend, timeout=600)
        print(f"  HTTP status: {response2.status_code}")
        if not print_result(response2.json()):
            sys.exit(1)
    except Exception as e:
        print(f"  Backend request error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
