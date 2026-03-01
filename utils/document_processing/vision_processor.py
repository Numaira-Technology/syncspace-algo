"""
Vision-Based Document Processor

Uses vision AI to understand Excel tables and update Word documents intelligently.
"""

import logging
import time
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from config import Config
from funnels.extract import read_docx
from funnels.llm_provider import get_llm_provider
from utils.excel_to_image import convert_excel_to_image
from utils.vision_data_extractor import extract_data_from_excel_image
from tqdm import tqdm
import json

logger = logging.getLogger(__name__)


def analyze_and_update_document(docx_path: str, excel_image_base64: str, 
                                extracted_data: Dict[str, Any], 
                                timeout: int = None) -> List[Tuple[str, str, float, str]]:
    """
    Analyze Word document and determine what needs to be updated based on Excel data.
    
    Args:
        docx_path: Path to Word document
        excel_image_base64: Base64 encoded Excel image
        extracted_data: Structured data extracted from Excel
        timeout: Optional timeout for LLM calls
        
    Returns:
        List of tuples: (original_sentence, updated_sentence, confidence, description)
    """
    logger.info("Starting vision-based document analysis")
    
    # Get LLM provider
    config = Config()
    llm_config = config.get_model_config()["llm"]
    timeout = timeout or llm_config.get("timeout", 30)
    llm = get_llm_provider()
    
    # Read Word document
    logger.info(f"Reading Word document: {docx_path}")
    sentences = read_docx(docx_path)
    logger.info(f"Found {len(sentences)} sentences in document")
    
    if not sentences:
        logger.warning("No sentences found in document")
        return []
    
    # Prepare extracted data summary for context
    metrics_list = list(extracted_data.get("metrics", {}).keys())
    time_periods = extracted_data.get("time_periods", [])
    
    logger.info(f"Available metrics: {metrics_list}")
    logger.info(f"Available time periods: {time_periods}")
    
    # Process each sentence to find ones that need updating
    results = []
    total_sentences = len(sentences)
    skipped_short = 0
    llm_call_count = 0
    llm_total_time = 0.0
    analysis_call_times = []
    update_call_times = []

    print("\nAnalyzing document with vision AI...")
    logger.info(f"Processing {total_sentences} sentences (2 LLM calls each if update needed)")
    
    for i, sentence in enumerate(tqdm(sentences.values(), desc="Analyzing", unit="sentence")):
        if len(sentence.strip()) < 20:
            skipped_short += 1
            logger.debug(f"[{i+1}/{total_sentences}] Skipping short sentence: '{sentence[:30]}'")
            continue
        
        logger.info(f"[{i+1}/{total_sentences}] Analyzing: '{sentence[:80]}...'")

        # Ask AI to analyze if this sentence contains financial data that should be updated
        analysis_prompt = f"""Analyze this sentence from a financial document:

Sentence: "{sentence}"

Available financial data from Excel:
- Metrics: {', '.join(metrics_list)}
- Time periods: {', '.join(time_periods)}

Task: Determine if this sentence contains financial data that should be updated with values from the Excel data.

Output a JSON object with this format:
{{
  "should_update": true/false,
  "reason": "Brief explanation",
  "metrics_mentioned": ["list", "of", "metrics", "if", "any"],
  "time_periods_mentioned": ["list", "of", "periods", "if", "any"]
}}

Only output the JSON, nothing else."""

        try:
            t_analysis = time.time()
            analysis_result = llm.analyze_text(analysis_prompt, timeout=timeout, max_tokens=500)
            analysis_elapsed = time.time() - t_analysis
            llm_call_count += 1
            llm_total_time += analysis_elapsed
            analysis_call_times.append(analysis_elapsed)
            logger.info(f"  [LLM call #{llm_call_count}] Analysis call: {analysis_elapsed:.2f}s")

            analysis_text = analysis_result.get("analysis", "")
            
            if not analysis_text or analysis_text == "none":
                logger.debug(f"  No analysis returned for sentence")
                continue
            
            # Parse analysis
            analysis_text = analysis_text.strip()
            if analysis_text.startswith("```json"):
                analysis_text = analysis_text[7:]
            elif analysis_text.startswith("```"):
                analysis_text = analysis_text[3:]
            if analysis_text.endswith("```"):
                analysis_text = analysis_text[:-3]
            analysis_text = analysis_text.strip()
            
            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError:
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group(0))
                else:
                    logger.debug(f"  Could not parse analysis JSON: {analysis_text[:100]}")
                    continue
            
            if not analysis.get("should_update", False):
                logger.debug(f"  No update needed — {analysis.get('reason', 'no reason given')}")
                continue
            
            logger.info(f"  Needs update — metrics: {analysis.get('metrics_mentioned', [])}, periods: {analysis.get('time_periods_mentioned', [])}")
            
            # Now ask AI to update the sentence with the correct data
            update_prompt = f"""Update this sentence with the correct financial data from the Excel table.

Original Sentence: "{sentence}"

Excel Data (JSON format):
{json.dumps(extracted_data, indent=2)}

Instructions:
1. Identify which metrics and time periods are mentioned in the sentence
2. Look up the correct values from the Excel data
3. Update ONLY the numerical values in the sentence
4. Keep the exact same sentence structure and wording
5. Maintain the same number formatting style (e.g., if it says "$X.XX billion", keep that format)
6. Do NOT change any other text

Output ONLY the updated sentence, nothing else. If you cannot update it confidently, output: SKIP"""

            t_update = time.time()
            update_result = llm.analyze_text(update_prompt, timeout=timeout, max_tokens=500)
            update_elapsed = time.time() - t_update
            llm_call_count += 1
            llm_total_time += update_elapsed
            update_call_times.append(update_elapsed)
            logger.info(f"  [LLM call #{llm_call_count}] Update call: {update_elapsed:.2f}s")

            updated_text = update_result.get("analysis", "").strip()
            
            if not updated_text or updated_text == "none" or updated_text == "SKIP":
                logger.warning(f"  Update call returned no result")
                continue
            
            if updated_text == sentence:
                logger.debug(f"  Update identical to original, skipping")
                continue
            
            len_ratio = len(updated_text) / len(sentence) if len(sentence) > 0 else 0
            len_similarity = 1.0 - abs(1.0 - len_ratio)
            
            metrics_present = sum(1 for m in analysis.get('metrics_mentioned', []) 
                                 if m.lower() in updated_text.lower())
            metric_score = metrics_present / max(len(analysis.get('metrics_mentioned', [])), 1)
            
            confidence = 0.6 + (len_similarity * 0.2) + (metric_score * 0.2)
            confidence = min(confidence, 0.99)
            
            description = f"Updated {', '.join(analysis.get('metrics_mentioned', []))} for {', '.join(analysis.get('time_periods_mentioned', []))}"
            
            logger.info(f"  Applied update (confidence: {confidence:.2f})")
            logger.info(f"    Original: '{sentence[:80]}'")
            logger.info(f"    Updated:  '{updated_text[:80]}'")
            
            results.append((sentence, updated_text, confidence, description))
            
        except Exception as e:
            logger.error(f"  Error processing sentence: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())
            continue

    # Summary stats
    logger.info("-" * 60)
    logger.info(f"SENTENCE ANALYSIS SUMMARY")
    logger.info(f"  Total sentences:     {total_sentences}")
    logger.info(f"  Skipped (too short): {skipped_short}")
    logger.info(f"  Processed:           {total_sentences - skipped_short}")
    logger.info(f"  Updates found:       {len(results)}")
    logger.info(f"  Total LLM calls:     {llm_call_count}")
    logger.info(f"  Total LLM time:      {llm_total_time:.2f}s")
    if analysis_call_times:
        logger.info(f"  Analysis calls — avg: {sum(analysis_call_times)/len(analysis_call_times):.2f}s, min: {min(analysis_call_times):.2f}s, max: {max(analysis_call_times):.2f}s")
    if update_call_times:
        logger.info(f"  Update calls   — avg: {sum(update_call_times)/len(update_call_times):.2f}s, min: {min(update_call_times):.2f}s, max: {max(update_call_times):.2f}s")
    logger.info("-" * 60)

    return results


def process_files_with_vision(docx_path: str, excel_path: str, timeout: int = None) -> List[Tuple[str, str, float, str]]:
    """
    Process Word and Excel files using vision-based approach.
    
    Args:
        docx_path: Path to Word document
        excel_path: Path to Excel file
        timeout: Optional timeout for LLM calls
        
    Returns:
        List of tuples: (original_sentence, updated_sentence, confidence, description)
    """
    logger.info("=" * 80)
    logger.info("VISION-BASED DOCUMENT PROCESSING")
    logger.info("=" * 80)
    pipeline_start = time.time()
    
    try:
        # Step 1: Convert Excel to image
        logger.info("Step 1: Converting Excel to image...")
        t0 = time.time()
        img, img_base64 = convert_excel_to_image(excel_path)
        logger.info(f"Step 1 done in {time.time() - t0:.2f}s — image size: {img.width}x{img.height}px, base64 length: {len(img_base64)} chars")
        
        # Step 2: Extract structured data from image
        logger.info("Step 2: Extracting data from Excel image using vision AI...")
        t0 = time.time()
        extracted_data = extract_data_from_excel_image(img_base64, timeout=timeout)
        logger.info(f"Step 2 done in {time.time() - t0:.2f}s — extracted {len(extracted_data.get('metrics', {}))} metrics")
        
        if not extracted_data.get("metrics"):
            logger.error("Failed to extract any metrics from Excel image")
            return []
        
        logger.info(f"Metrics: {list(extracted_data['metrics'].keys())}")
        
        # Step 3: Analyze Word document and determine updates
        logger.info("Step 3: Analyzing Word document and determining updates...")
        t0 = time.time()
        results = analyze_and_update_document(docx_path, img_base64, extracted_data, timeout)
        logger.info(f"Step 3 done in {time.time() - t0:.2f}s — {len(results)} updates identified")
        
        total = time.time() - pipeline_start
        logger.info("=" * 80)
        logger.info(f"PROCESSING COMPLETE: {len(results)} updates in {total:.2f}s total")
        logger.info("=" * 80)
        
        return results
        
    except Exception as e:
        logger.error(f"Error in vision-based processing after {time.time() - pipeline_start:.2f}s: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return []


def analyze_and_update_document_from_json(
        word_sentences: Dict[str, str],
        excel_rows: List[Dict[str, Any]],
        timeout: int = None) -> List[Tuple[str, str, float, str]]:
    """
    Analyze pre-extracted sentences and Excel rows (from frontend) and determine updates.

    Args:
        word_sentences: Dict of { key: sentence_text } from the Word document
        excel_rows: List of { row_header, values: { period: value } } from the Excel file
        timeout: Optional timeout for LLM calls

    Returns:
        List of tuples: (original_sentence, updated_sentence, confidence, description)
    """
    logger.info("Starting JSON-based document analysis")

    config = Config()
    llm_config = config.get_model_config()["llm"]
    timeout = timeout or llm_config.get("timeout", 30)
    llm = get_llm_provider()

    # Build a structured metrics dict from excel rows for LLM context
    metrics_dict: Dict[str, Any] = {}
    for row in excel_rows:
        if not isinstance(row, dict):
            continue
        row_header = row.get("row_header") or row.get("row_header", "")
        values = row.get("values", {})
        if row_header and values:
            metrics_dict[str(row_header)] = values

    metrics_list = list(metrics_dict.keys())
    time_periods = list({period for vals in metrics_dict.values() if isinstance(vals, dict) for period in vals.keys()})

    logger.info(f"Available metrics: {metrics_list}")
    logger.info(f"Available time periods: {time_periods}")

    sentences = word_sentences
    total_sentences = len(sentences)
    logger.info(f"Found {total_sentences} sentences in document")

    if not sentences:
        logger.warning("No sentences found in document")
        return []

    results = []
    skipped_short = 0
    llm_call_count = 0
    llm_total_time = 0.0
    analysis_call_times = []
    update_call_times = []

    logger.info(f"Processing {total_sentences} sentences (2 LLM calls each if update needed)")

    for i, sentence in enumerate(sentences.values()):
        if not isinstance(sentence, str):
            sentence = str(sentence)
        if len(sentence.strip()) < 20:
            skipped_short += 1
            continue

        logger.info(f"[{i+1}/{total_sentences}] Analyzing: '{sentence[:80]}...'")

        # Determine the most recent time period so the update prompt can use it
        most_recent_periods = sorted(time_periods, reverse=True) if time_periods else []
        most_recent_label = most_recent_periods[0] if most_recent_periods else "most recent period"

        analysis_prompt = f"""You are reviewing a financial document to find sentences that contain specific dollar amounts, percentages, or numerical counts for financial metrics.

Sentence from the document:
"{sentence}"

Financial metrics available in the Excel data:
{', '.join(metrics_list)}

Task: Does this sentence contain ANY specific numerical values (dollar amounts, percentages, unit counts) that correspond to one or more of the above financial metrics?

Output a JSON object:
{{
  "should_update": true/false,
  "reason": "Brief explanation",
  "metrics_mentioned": ["metrics referenced in this sentence"],
  "time_periods_mentioned": ["time periods referenced in this sentence"]
}}

Set should_update to TRUE if the sentence contains specific dollar amounts, percentages, or counts for ANY of the listed financial metrics — even if the values look correct.
Set should_update to FALSE only if the sentence contains NO specific financial figures at all.
Output ONLY the JSON, nothing else."""

        try:
            t_analysis = time.time()
            analysis_result = llm.analyze_text(analysis_prompt, timeout=timeout, max_tokens=500)
            analysis_elapsed = time.time() - t_analysis
            llm_call_count += 1
            llm_total_time += analysis_elapsed
            analysis_call_times.append(analysis_elapsed)
            logger.info(f"  [LLM call #{llm_call_count}] Analysis call: {analysis_elapsed:.2f}s")

            analysis_text = analysis_result.get("analysis", "")
            if not analysis_text or analysis_text == "none":
                continue

            analysis_text = analysis_text.strip()
            if analysis_text.startswith("```json"):
                analysis_text = analysis_text[7:]
            elif analysis_text.startswith("```"):
                analysis_text = analysis_text[3:]
            if analysis_text.endswith("```"):
                analysis_text = analysis_text[:-3]
            analysis_text = analysis_text.strip()

            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError:
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group(0))
                else:
                    logger.debug(f"  Could not parse analysis JSON: {analysis_text[:100]}")
                    continue

            if not analysis.get("should_update", False):
                logger.debug(f"  No update needed — {analysis.get('reason', 'no reason given')}")
                continue

            logger.info(f"  Needs update — metrics: {analysis.get('metrics_mentioned', [])}, periods: {analysis.get('time_periods_mentioned', [])}")

            update_prompt = f"""You are updating a financial document sentence to reflect the most recent data from the Excel spreadsheet.

Original Sentence:
"{sentence}"

Excel Data (all available time periods and values):
{json.dumps(metrics_dict, indent=2)}

Most recent time period in the Excel: {most_recent_label}

Instructions:
1. Identify each financial metric and the time period mentioned in the sentence.
2. For each value, find the EQUIVALENT value for the MOST RECENT equivalent time period in the Excel.
   Example: if the sentence says "three months ended June 30, 2023", look for "Three Months Ended June 30, 2024" data.
3. Replace BOTH the old time period reference AND the corresponding numerical values with the updated ones.
4. Keep the exact same sentence structure and wording otherwise.
5. Use the same number formatting style (billions, millions, percentages, etc.).
6. Do NOT change any text that is not a financial figure or time period.

Output ONLY the updated sentence. If you cannot update it confidently, output: SKIP"""

            t_update = time.time()
            update_result = llm.analyze_text(update_prompt, timeout=timeout, max_tokens=500)
            update_elapsed = time.time() - t_update
            llm_call_count += 1
            llm_total_time += update_elapsed
            update_call_times.append(update_elapsed)
            logger.info(f"  [LLM call #{llm_call_count}] Update call: {update_elapsed:.2f}s")

            updated_text = update_result.get("analysis", "").strip()

            if not updated_text or updated_text == "none" or updated_text == "SKIP":
                logger.warning(f"  Update call returned no result")
                continue

            if updated_text == sentence:
                logger.debug(f"  Update identical to original, skipping")
                continue

            len_ratio = len(updated_text) / len(sentence) if len(sentence) > 0 else 0
            len_similarity = 1.0 - abs(1.0 - len_ratio)

            metrics_present = sum(1 for m in analysis.get('metrics_mentioned', [])
                                 if m.lower() in updated_text.lower())
            metric_score = metrics_present / max(len(analysis.get('metrics_mentioned', [])), 1)

            confidence = 0.6 + (len_similarity * 0.2) + (metric_score * 0.2)
            confidence = min(confidence, 0.99)

            description = f"Updated {', '.join(analysis.get('metrics_mentioned', []))} for {', '.join(analysis.get('time_periods_mentioned', []))}"

            logger.info(f"  Applied update (confidence: {confidence:.2f})")
            results.append((sentence, updated_text, confidence, description))

        except Exception as e:
            logger.error(f"  Error processing sentence: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())
            continue

    logger.info("-" * 60)
    logger.info(f"JSON ANALYSIS SUMMARY")
    logger.info(f"  Total sentences:     {total_sentences}")
    logger.info(f"  Skipped (too short): {skipped_short}")
    logger.info(f"  Updates found:       {len(results)}")
    logger.info(f"  Total LLM calls:     {llm_call_count}")
    logger.info(f"  Total LLM time:      {llm_total_time:.2f}s")
    logger.info("-" * 60)

    return results


def analyze_and_update_from_images(
        docx_image_b64: str,
        excel_image_b64: str,
        timeout: int = None) -> List[Tuple[str, str, float, str]]:
    """
    Process Word document and Excel file, both provided as base64 PNG images.

    Step 1 — Excel vision OCR: use VisionDataExtractor to extract structured metrics.
    Step 2 — Word vision OCR: call vision LLM to extract clean meaningful sentences
              (paragraphs, table cells), excluding decorative headers, page numbers, footers.
    Step 3 — Text LLM analysis: call analyze_and_update_document_from_json to find and
              produce updates using the structured Excel data.

    Args:
        docx_image_b64: Base64-encoded PNG of the rendered Word document
        excel_image_b64: Base64-encoded PNG of the rendered Excel sheet
        timeout: Optional timeout for LLM calls

    Returns:
        List of tuples: (original_sentence, updated_sentence, confidence, description)
    """
    logger.info("=" * 80)
    logger.info("IMAGE-BASED DOCUMENT PROCESSING")
    logger.info("=" * 80)
    pipeline_start = time.time()

    config = Config()
    llm_config = config.get_model_config()["llm"]
    timeout = timeout or llm_config.get("timeout", 60)
    llm = get_llm_provider()

    # Step 1: Extract structured financial data from Excel image
    logger.info("Step 1: Extracting financial data from Excel image via vision OCR...")
    t0 = time.time()
    from utils.vision_data_extractor import VisionDataExtractor
    extractor = VisionDataExtractor(llm_provider=llm)
    extracted_data = extractor.extract_financial_data(excel_image_b64, timeout=timeout)
    logger.info(f"Step 1 done in {time.time() - t0:.2f}s — extracted {len(extracted_data.get('metrics', {}))} metrics")

    if not extracted_data.get("metrics"):
        logger.error("Failed to extract any metrics from Excel image")
        return []

    logger.info(f"Metrics: {list(extracted_data['metrics'].keys())}")

    # Step 2: OCR the Word document image to extract clean sentences
    logger.info("Step 2: OCR Word document image to extract clean content...")
    t0 = time.time()

    word_ocr_prompt = """You are reading an image of a Word document. 
Extract ONLY meaningful body content as individual sentences or entries:
- Full paragraph sentences
- Table cell text as plain sentences
- Captions and labels with their associated data

EXCLUDE entirely:
- Page numbers (e.g. "1", "Page 1 of 5")
- Decorative section headers (e.g. "ITEM 2.", "Overview", "Introduction")
- Running headers and footers
- Watermarks
- Blank lines

Return a JSON object in this exact format:
{ "sentences": ["sentence 1", "sentence 2", ...] }

Each entry must be a complete, standalone piece of meaningful content (at least 5 words).
Output ONLY the JSON, nothing else."""

    word_ocr_result = llm.analyze_image(word_ocr_prompt, docx_image_b64, timeout=timeout, max_tokens=4000)
    word_ocr_text = word_ocr_result.get("analysis", "")
    logger.info(f"Step 2 done in {time.time() - t0:.2f}s — raw OCR length: {len(word_ocr_text)}")

    # Parse the OCR response
    sentences_list: List[str] = []
    if word_ocr_text and word_ocr_text != "none":
        ocr_text = word_ocr_text.strip()
        if ocr_text.startswith("```json"):
            ocr_text = ocr_text[7:]
        elif ocr_text.startswith("```"):
            ocr_text = ocr_text[3:]
        if ocr_text.endswith("```"):
            ocr_text = ocr_text[:-3]
        ocr_text = ocr_text.strip()

        try:
            ocr_data = json.loads(ocr_text)
            sentences_list = ocr_data.get("sentences", [])
        except json.JSONDecodeError:
            import re
            json_match = re.search(r'\{.*\}', ocr_text, re.DOTALL)
            if json_match:
                try:
                    ocr_data = json.loads(json_match.group(0))
                    sentences_list = ocr_data.get("sentences", [])
                except json.JSONDecodeError:
                    logger.error("Could not parse Word OCR JSON response")
            else:
                logger.error(f"No JSON found in Word OCR response: {ocr_text[:200]}")

    if not sentences_list:
        logger.error("No sentences extracted from Word document image")
        return []

    logger.info(f"Extracted {len(sentences_list)} sentences from Word document")

    # Convert list to keyed dict for analyze_and_update_document_from_json
    sentences_dict = {f"s{i}": s for i, s in enumerate(sentences_list)}

    # Convert extracted metrics into the excel_rows format
    excel_rows = [
        {"row_header": metric_name, "values": period_values}
        for metric_name, period_values in extracted_data["metrics"].items()
    ]

    # Step 3: Run text-LLM analysis to find sentences needing updates
    logger.info("Step 3: Running text-LLM analysis to identify and generate updates...")
    t0 = time.time()
    results = analyze_and_update_document_from_json(sentences_dict, excel_rows, timeout=timeout)
    logger.info(f"Step 3 done in {time.time() - t0:.2f}s — {len(results)} updates identified")

    total = time.time() - pipeline_start
    logger.info("=" * 80)
    logger.info(f"IMAGE PROCESSING COMPLETE: {len(results)} updates in {total:.2f}s total")
    logger.info("=" * 80)

    return results


def apply_updates_to_document(docx_path: str, output_path: str, 
                              updates: List[Tuple[str, str, float, str]],
                              min_confidence: float = 0.7) -> int:
    """
    Apply updates to Word document with highlighting.
    
    Args:
        docx_path: Path to original Word document
        output_path: Path to save updated document
        updates: List of (original, updated, confidence, description) tuples
        min_confidence: Minimum confidence to apply update
        
    Returns:
        Number of updates applied
    """
    logger.info(f"Applying updates to document (min confidence: {min_confidence})")
    
    doc = Document(docx_path)
    updates_applied = 0
    
    # Create a mapping of original to updated text
    update_map = {orig: (upd, conf, desc) for orig, upd, conf, desc in updates if conf >= min_confidence}
    
    logger.info(f"Total updates to apply: {len(update_map)}")
    
    for paragraph in doc.paragraphs:
        para_text = paragraph.text
        
        # Check if this paragraph contains any text to update
        for original_text, (updated_text, confidence, description) in update_map.items():
            if original_text in para_text:
                logger.info(f"Applying update in paragraph (confidence: {confidence:.2f})")
                logger.debug(f"  Original: '{original_text[:80]}...'")
                logger.debug(f"  Updated:  '{updated_text[:80]}...'")
                
                # Replace the text
                para_text = para_text.replace(original_text, updated_text)
                
                # Clear existing runs and create new one with highlighting
                paragraph.clear()
                run = paragraph.add_run(para_text)
                
                # Highlight the updated text in yellow
                run.font.highlight_color = 7  # Yellow
                
                updates_applied += 1
                break  # Only apply one update per paragraph
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Document saved to: {output_path}")
    logger.info(f"Total updates applied: {updates_applied}")
    
    return updates_applied

